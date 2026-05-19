"""
Generate the Appendix section (from page 87 onwards) for the LMS Report.
Uses python-docx to create a professional academic-style DOCX document.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Screenshot directory ──────────────────────────────────────────────
SCREENSHOT_DIR = r"C:\Users\nagra\.gemini\antigravity\brain\0ef93c43-96fb-45ec-a4b6-4f40aedb8f69"
OUTPUT_DIR = r"d:\LIBRARY_MANAGEMENT_SYSTEM"

# ── Ordered screenshot groups (module-wise) ───────────────────────────
APPENDIX_SECTIONS = [
    {
        "title": "A.1  Authentication Module",
        "screenshots": [
            ("login_portal_1779181656102.png", "Login Portal – Role Selection Page"),
            ("student_login_1779181688867.png", "Student Login Page"),
            ("student_registration_1779181786109.png", "Student Registration Page"),
            ("admin_login_1779182274682.png", "Administrator Login Page"),
            ("librarian_login_1779181844879.png", "Librarian Login Page"),
            ("forgot_password_1779181871912.png", "Forgot Password Page"),
        ]
    },
    {
        "title": "A.2  Home Page",
        "screenshots": [
            ("home_page_hero_1779181503295.png", "Home Page – Hero Section with Navigation Bar"),
            ("home_page_categories_1779181516241.png", "Home Page – Browse by Categories Section"),
            ("home_page_new_arrivals_1779181548067.png", "Home Page – New Arrivals Section"),
            ("home_page_footer_1779181564861.png", "Home Page – Library Hours and Footer Section"),
        ]
    },
    {
        "title": "A.3  Books Catalogue",
        "screenshots": [
            ("books_listing_1779182104735.png", "All Books Listing with Category Sidebar and Search"),
        ]
    },
    {
        "title": "A.4  Categories Page",
        "screenshots": [
            ("categories_1779182129913.png", "Browse All Categories Page"),
        ]
    },
    {
        "title": "A.5  About Us Page",
        "screenshots": [
            ("about_us_top_1779181962295.png", "About Us Page – Header and Mission Section"),
            ("about_us_bottom_1779182025142.png", "About Us Page – Services and Team Section"),
        ]
    },
    {
        "title": "A.6  Contact Us Page",
        "screenshots": [
            ("contact_us_1779182079718.png", "Contact Us Page – Information Cards and Message Form"),
        ]
    },
    {
        "title": "A.7  Admin Dashboard",
        "screenshots": [
            ("admin_dashboard_1779182359878.png", "Admin Dashboard – Overview with Statistics, Category Mix Chart, and New Acquisitions"),
        ]
    },
    {
        "title": "A.8  Add Book Module",
        "screenshots": [
            ("add_book_form_1779182387237.png", "Add New Acquisition Form – Basic Information and Publication Details"),
        ]
    },
    {
        "title": "A.9  Inventory Management",
        "screenshots": [
            ("library_inventory_1779182422888.png", "Library Inventory – Book Cards with Edit/Delete Controls"),
        ]
    },
    {
        "title": "A.10  Member Management",
        "screenshots": [
            ("member_management_1779182448603.png", "Library Members Directory – Member List with Status and Actions"),
        ]
    },
    {
        "title": "A.11  Librarian / Staff Management",
        "screenshots": [
            ("add_librarian_1779182755974.png", "Add New Librarian – Staff Registration Form"),
        ]
    },
    {
        "title": "A.12  Book Issue and Return",
        "screenshots": [
            ("issue_requests_1779182784392.png", "Pending Book Issue Requests Page"),
            ("return_requests_1779182814856.png", "Return Book Requests Page"),
            ("issued_books_1779182847790.png", "Books Issued – Currently Borrowed Books List"),
        ]
    },
    {
        "title": "A.13  Reservation Management",
        "screenshots": [
            ("admin_reservations_1779183046637.png", "Reservation Queue – Status Cards and Filter Tabs"),
        ]
    },
    {
        "title": "A.14  Reports and Analytics",
        "screenshots": [
            ("admin_reports_1779183070345.png", "Reports Module – System Summary with Export Option"),
        ]
    },
    {
        "title": "A.15  Fine Management",
        "screenshots": [
            ("fine_management_1779183161789.png", "Fine Management – Outstanding Fines Dashboard"),
            ("fine_configuration_1779183186778.png", "Fine Configuration – Rules and Settings"),
        ]
    },
    {
        "title": "A.16  Responsive / Mobile Views",
        "screenshots": [
            ("mobile_home_page_1779183330997.png", "Mobile Responsive View – Home Page"),
            ("mobile_books_page_1779183286992.png", "Mobile Responsive View – Books Catalogue"),
            ("mobile_admin_dashboard_1779183263572.png", "Mobile Responsive View – Admin Dashboard"),
        ]
    },
]


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_page_number(doc):
    """Add page number to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = " PAGE "
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)

        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'


def set_page_start_number(doc, start_num):
    """Set starting page number."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), str(start_num))
        sectPr.append(pgNumType)


def create_appendix():
    doc = Document()
    
    # ── Page Setup ────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(2.54)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # ── Default font ──────────────────────────────────────────────────
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # ── Heading styles ────────────────────────────────────────────────
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_rpr = h1.element.get_or_add_rPr()
    h1_rFonts = OxmlElement('w:rFonts')
    h1_rFonts.set(qn('w:ascii'), 'Times New Roman')
    h1_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    h1_rpr.insert(0, h1_rFonts)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2_rpr = h2.element.get_or_add_rPr()
    h2_rFonts = OxmlElement('w:rFonts')
    h2_rFonts.set(qn('w:ascii'), 'Times New Roman')
    h2_rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    h2_rpr.insert(0, h2_rFonts)

    # ── Set starting page number to 87 ────────────────────────────────
    set_page_start_number(doc, 87)
    add_page_number(doc)

    # ── APPENDIX TITLE PAGE ───────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("APPENDIX")
    run.bold = True
    run.font.size = Pt(28)
    run.font.name = 'Times New Roman'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Application Screenshots")
    run2.font.size = Pt(16)
    run2.font.name = 'Times New Roman'
    run2.font.color.rgb = RGBColor(80, 80, 80)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_before = Pt(24)
    run3 = desc.add_run(
        "This appendix contains comprehensive screenshots of the Library Management System\n"
        "application, organized module-wise. These screenshots demonstrate the user interface,\n"
        "functionality, and responsive design of the developed system."
    )
    run3.font.size = Pt(12)
    run3.font.name = 'Times New Roman'
    run3.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()

    # ── APPENDIX CONTENT ──────────────────────────────────────────────
    figure_counter = 1

    for section_data in APPENDIX_SECTIONS:
        # Section heading
        heading = doc.add_heading(section_data["title"], level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for filename, caption in section_data["screenshots"]:
            filepath = os.path.join(SCREENSHOT_DIR, filename)
            
            if not os.path.exists(filepath):
                print(f"  [!] Missing: {filename}")
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[Screenshot not available: {filename}]")
                run.font.color.rgb = RGBColor(200, 0, 0)
                run.font.size = Pt(10)
                figure_counter += 1
                continue

            # Add spacing before image
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(6)
            spacer.paragraph_format.space_after = Pt(0)

            # Add screenshot
            img_paragraph = doc.add_paragraph()
            img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_paragraph.paragraph_format.space_before = Pt(0)
            img_paragraph.paragraph_format.space_after = Pt(4)

            # Determine image width based on type
            if "mobile" in filename.lower():
                img_width = Inches(2.8)
            else:
                img_width = Inches(5.8)

            run = img_paragraph.add_run()
            run.add_picture(filepath, width=img_width)

            # Add figure caption
            caption_para = doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_para.paragraph_format.space_before = Pt(4)
            caption_para.paragraph_format.space_after = Pt(12)

            cap_run = caption_para.add_run(f"Figure A.{figure_counter}: {caption}")
            cap_run.bold = True
            cap_run.font.size = Pt(10)
            cap_run.font.name = 'Times New Roman'
            cap_run.font.color.rgb = RGBColor(50, 50, 50)

            print(f"  [OK] Added Figure A.{figure_counter}: {caption}")
            figure_counter += 1

        # Add some spacing between sections
        doc.add_paragraph("")

    # ── FINAL SUMMARY TABLE ───────────────────────────────────────────
    doc.add_page_break()
    summary_heading = doc.add_heading("A.17  Summary of Application Screenshots", level=2)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    run = intro.add_run(
        "Table A.1 provides a consolidated summary of all application screenshots "
        "included in this appendix, organized by module."
    )
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

    # Create summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header row
    hdr = table.rows[0]
    headers = ['S.No.', 'Module', 'No. of Screenshots']
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Times New Roman'
        # Shade header
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D9E2F3')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    total_screenshots = 0
    for idx, section_data in enumerate(APPENDIX_SECTIONS, 1):
        row = table.add_row()
        count = len(section_data["screenshots"])
        total_screenshots += count
        
        values = [str(idx), section_data["title"].split("  ", 1)[1], str(count)]
        for i, val in enumerate(values):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 1 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Times New Roman'

    # Total row
    total_row = table.add_row()
    total_row.cells[0].text = ''
    p = total_row.cells[1].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Total")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    p2 = total_row.cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(str(total_screenshots))
    run2.bold = True
    run2.font.size = Pt(11)
    run2.font.name = 'Times New Roman'
    
    for cell in total_row.cells:
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'E2EFDA')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    # Table caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    run = cap.add_run("Table A.1: Summary of Application Screenshots by Module")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.8)
        row.cells[1].width = Inches(3.5)
        row.cells[2].width = Inches(1.7)

    # ── Save DOCX ─────────────────────────────────────────────────────
    docx_path = os.path.join(OUTPUT_DIR, "LMS_Report_Appendix.docx")
    doc.save(docx_path)
    print(f"\n[DONE] DOCX saved: {docx_path}")
    return docx_path


if __name__ == "__main__":
    create_appendix()
